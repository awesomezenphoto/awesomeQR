'''
Functions of awe-qr.py
1. input file for names. one name per line.
2. generate qr code for each name. 
3. save qr image to a word document. 'myQRStudents.docx'
ex. awesome-qr.exe <names-file> 
the qr code will be save in myQRStudents.docx file.....

For development requirement:
1. python 3.5
2. python qr code
3. python docx

For Release:
run > pyinstaller awesome-qr to generate a dist folder.
which can be used by anyone without installing python and modules. 
'''
import qrcode
import sys
import string
from docx import Document
from docx.shared import Inches 

Revision = "0.1"
MAXCOLS = 5
MAXROW = 4
SAVEFILE = 'myQRStudents.docx'

try:
    in_file = open(sys.argv[1], "r")
except:
    sys.exit("ERROR. File is NOT Found from Your Input FileName, Please Try Again")
	
#read input file, and keep in lines 
lines = in_file.readlines()
in_file.close()
# create a document
doc = Document()
#paragraph = doc.add_paragraph("===   AWESOME   ===\r\n")

stcount = 0
rowcount = 0
table = doc.add_table(rows=0, cols=MAXCOLS)


for line in lines:
	#print (line)
	name = ""
	name=" ".join(line.split())
	#name += "\n"
	img = qrcode.make(name)
	img.save('tmp.png')
		
	if (stcount % MAXCOLS) == 0 :
		if (rowcount > MAXROW ) :
			rowcount = 0
			doc.add_page_break()
			print ("========================== Next Page, ", stcount)
			table = doc.add_table(rows=0, cols=MAXCOLS)
		row_cells = table.add_row().cells 
		rowcount += 1
		print ("-------------new row", rowcount)
		
	print (name)
	paragraph=row_cells[stcount%MAXCOLS].paragraphs[0]
	run = paragraph.add_run()
	run.add_text(line.strip())
	run.add_picture('tmp.png', width=Inches(1.00))
	stcount += 1

	##doc.add_paragraph(line.strip())
	##doc.add_picture('tmp.png', width=Inches(1.00))
	#img.save(line.strip()+'.png')
	
doc.save(SAVEFILE)
print('*********************************** Rev:', Revision )
print("Created:", stcount, "qr-codes and saved as", SAVEFILE )	